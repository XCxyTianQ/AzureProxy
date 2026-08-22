# -*- coding: utf-8 -*-
"""Generate AzureProxy first technical report docx (AzureDoc series).

Style mirror of AzureBranches-26.1.2-EXP7.docx:
  - title block centered (project 22pt bold / title 16pt bold / subtitle 12pt / meta 10.5pt)
  - 摘要 + 关键词 justified 10.5pt (inline bold for key terms)
  - h1 = 14pt bold left, h2 = 12pt bold left, body justified 10.5pt
  - tables: Table Grid, header row bold
Output: F:\\AzureCore\\AzureDoc\\AzureProxy-26.1.2-0001.docx
"""
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = r'F:\AzureCore\AzureDoc\AzureProxy-26.1.2-0001.docx'
FONT = 'Times New Roman'

doc = docx.Document()
st = doc.styles['Normal']
st.font.name = FONT
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def mkruns(p, text, bold=False):
    for chunk in text.split('**'):
        if not chunk:
            continue
        r = p.add_run(chunk)
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        r.font.size = Pt(10.5)
        r.bold = bold


def para(text, size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space=Pt(3)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space
    p.paragraph_format.line_spacing = 1.0
    mkruns(p, text, bold)
    for r in p.runs:
        r.font.size = Pt(size)
    return p


def h1(text):
    return para(text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space=Pt(8))


def h2(text):
    return para(text, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space=Pt(6))


def blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return p


def table(rows, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            c = t.cell(ri, ci)
            c.text = ''
            p = c.paragraphs[0]
            mkruns(p, cell, bold=(ri == 0))
            for r in p.runs:
                r.font.size = Pt(10.5)
    return t


# ---------------- title block ----------------
para('AzureProxy', size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(2))
para('Velocity 下游代理与 EXP7 代理侧配套（v1 工程报告）', size=16, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(2))
para('从「Velocity 4.1.0-SNAPSHOT 下游构建管线」到「26.1.2 协议对齐 · azureproxy.mode EXP 三档预设 · '
     '服务器切换课题 · T1 命令树注入修复」的完整工程记录', size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(4))
para('版本 4.1.0-SNAPSHOT（基线 dev/4.0.0 @ 4772ca3）　　2026 年 08 月 23 日',
     size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(2))
para('基于 Velocity 4.1.0-SNAPSHOT / Minecraft 26.1.2（protocol 775）',
     size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(2))
para('致谢 PaperMC / Velocity（Tux 及其贡献者）——本下游基于 dev/4.0.0 分支；'
     '配套 AzureBranches EXP7（b_linear_v4 存储引擎）', size=10.5,
     align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(6))
blank()

# ---------------- 摘要 ----------------
para('**摘要**')
para('**AzureProxy** 是 AzureCore 下与 AzureBranches 并列的 Velocity GPLv3 下游（基线 '
     'dev/4.0.0 @ 4772ca3，版本 4.1.0-SNAPSHOT），承担 EXP7 的代理侧配套：仓库只持有构建驱动、'
     '补丁与文档，上游源码按固定 ref 克隆到 build/velocity-src 后打补丁构建（HEAD 漂移即失败），'
     '构建链 cloneVelocity → applyAzurePatches → buildVelocity → buildAzureProxyJar（Gradle 9.4.1 '
     '复用 AzureBranches，shadowJar 产物 azureproxy-proxy-4.1.0-SNAPSHOT-all.jar，启动横幅 '
     'Booting up AzureProxy）。代理侧提供 **azureproxy.mode 三档预设**（SAFE/ACCESS/EXP）——在 '
     'nightconfig 绑定**之前**改写原始配置，预设经上游构造器/迁移/校验链自然生效；EXP 档强制 '
     '**announce-proxy-commands=true** 修复 **T1**（客户端命令树缺少代理命令 → /server 红色无补全）：'
     '真实客户端实测合并树 children 26→28、命令包 pid=0x10 len=718 含 velocity:callback 与 '
     'server 子树，/server tab 补全正常。**26.1.2 协议对齐**结论：Velocity 上游 26_1 表与真实 '
     '客户端一致（KeepAlive 0x2C / JoinGame 0x31 / SystemChat 0x79 / AvailableCommands 0x10），'
     'BundleDelimiter 占据 CB 注册表 0x00 造成的「索引抽取 +1」假象曾导致 27 处误报修复（已 revert），'
     '方法学转为真实客户端 + 合成客户端双向互证。**服务器切换课题**全部实测：/server 双后端切换、'
     '世界隔离（/say 只广播本世界）、每后端 ops.json 独立、现代转发 UUID 跨后端恒定 '
     '（2b47bbd5-…），back-end 宕机 fallback 同秒连接另一后端（无世界加载等待）。'
     '26.2（776）验证暂缓，列入后续 P0。')
para('**关键词：**Velocity；下游；4.1.0-SNAPSHOT；proxy；azureproxy.mode；SAFE/ACCESS/EXP；'
     'announce-proxy-commands；命令树注入；26.1.2；protocol 775；StateRegistry；现代转发；fallback；'
     '服务器切换；GPLv3')
blank()

# ---------------- 一 ----------------
h1('一、背景与定位')
h2('1.1　为什么需要代理侧下游')
para('AzureBranches EXP7（b_linear_v4 存储引擎）在 26.1.2 上验证了后端侧改造，但其形态是「直连后端」，'
     '缺少代理层：多后端编排（服务器切换 / fallback）、转发身份、代理命令面与后端命令面合并——这些'
     '由代理承担。Velocity 的 **dev/4.0.0** 是携带 26.1 协议支持的主线，但直接使用上游二进制无法：'
     '(1) 表达 EXP 三档配置预设（与 AzureBranches command_blocks.mode 呼应）；(2) pin 已验证基线、'
     '保证可复现构建；(3) 形成「代理-后端」配对课题（T1/转发/切换）。因此建立 AzureProxy 下游。')
h2('1.2　与 AzureBranches 的关系')
para('同属 AzureCore，互为配套：**AzureProxy 不改后端**——后端侧只要求 paper-global.yml 的 '
     'velocity.enabled=true（现代转发开关）与每后端独立 ops.json；代理侧全部改动只在 Velocity 的'
     '配置面与命令面。测试拓扑：proxy 25571 → exp7（25570 / RCON 25576，world-exp7-v4，'
     'b_linear_v4 引擎）与 arena（25572 / RCON 25578，seed 987654321）。')
h2('1.3　基线策略')
para('velocityRef 固定 dev/4.0.0 @ 4772ca3022c49bfab37c703f72cbca7654fb5848；cloneVelocity 对已有'
     '克隆做 HEAD 严格校验（漂移即抛错，提示 re-baseline）；提升上游按 README「版本同步」流程显式'
     '执行（bump ref → 删克隆 → 重克隆 → overlay fail-fast 逐项修复 → 全绿提交）。')

# ---------------- 二 ----------------
h1('二、总体架构')
h2('2.1　仓库边界')
para('仓库只持有驱动/补丁/文档/工具，上游克隆与构建产物全部 gitignored：')
table([
    ['组件', '职责'],
    ['build.gradle.kts', '构建驱动：pin ref / 补丁应用 / 品牌 / 打包（4 个 Gradle 任务）'],
    ['azurepatches-src/', '整文件覆盖层（必须有对应上游文件，fail-fast）'],
    ['azurepatches-new/', '新增类（com.azureproxy.*）'],
    ['gen-velocity-config-overlay.py', '由上游 HEAD 原文一键再生成覆盖层（唯一锚点插入）'],
    ['mcclient.py / mcping.py', '合成 protocol-775 客户端（E2E）/ status 探测'],
    ['TECHNICAL.md', '原理向技术文档（构建/补丁/预设/协议/验证矩阵）'],
    ['gradlew 9.4.1', '复用 AzureBranches 的 Gradle distribution（不下载上游 9.6.1）'],
])
h2('2.2　构建管线')
para('任务链：buildAzureProxyJar → buildVelocity → applyAzurePatches → cloneVelocity；关键决策：')
table([
    ['任务', '职责', '关键决策'],
    ['cloneVelocity', '固定 ref 克隆/校验（幂等）', 'fetch --depth 1 + checkout --detach；已有克隆 HEAD ≠ pin 即失败'],
    ['applyAzurePatches', '应用补丁', 'overlay 先全量校验上游存在（README 约定文件跳过）再覆盖；new 直接拷贝'],
    ['buildVelocity', '品牌 + 编译', '先 git checkout 还原 proxy/build.gradle.kts（transformSource 非幂等，靠还原保证幂等）→ 品牌 2 锚点 → wrapper pin 9.4.1 → compileJava'],
    ['buildAzureProxyJar', '打包', 'shadowJar → 挑最大非 -sources/-javadoc jar → velocity- 前缀重命名为 azureproxy- → 拷至 build/libs'],
])
para('**transformSource** 为 fail-fast 唯一锚点替换：锚点缺失或匹配 >1 次立即打断构建——上游漂移'
     '不会静默通过。启动横幅 Booting up AzureProxy 由 Manifest 的 Implementation-Title/Vendor 提供'
     '（VelocityServer.getVersion() 读取），实测日志确认该行。')
h2('2.3　补丁系统与覆盖层生成器')
para('VelocityConfiguration.java 覆盖层不是手敲的：脚本从 build/velocity-src 的 HEAD:… 读上游原文，'
     '对唯一锚点 PacketLimiterConfig.fromConfig(...) 做插入后写回 azurepatches-src。re-baseline '
     '后重跑一次即可再生成（锚点重复/缺失 assert 失败即失败）。约定见 azurepatches-src/new 各自 '
     'README（跳过逻辑、包前缀 com.azureproxy.*、风格随上游以通过 spotless）。')

# ---------------- 三 ----------------
h1('三、azureproxy.mode 预设体系')
h2('3.1　动机与配置')
para('与 AzureBranches 的 command_blocks.mode（SAFE/ACCESS/EXP）同构：一个开关把代理侧网络/命令面'
     '整体重调到目标后端族，而不是让用户手工改多个 Velocity 选项。velocity.toml：')
para('[azureproxy]\nmode = "EXP"        # SAFE（默认）| ACCESS | EXP', bold=False)
para('未知值 fail-soft：打印告警并保持 SAFE。')
h2('3.2　挂载点（为什么在这里）')
para('见 VelocityConfiguration.read() 覆盖层：在 PacketLimiterConfig 绑定**之后**、forwarding-secret '
     '非空校验**之前**调用 AzureProxyMode.applyToConfig(config, advancedConfig)。选型依据：')
para('· 在 nightconfig 绑定之前改写原始配置——预设经上游正常构造器/迁移/校验链生效，不绕过任何语义；')
para('· 在 forwarding-secret 校验之前——EXP 强制 MODERN 时，上游对 forwarding secret 的既有校验'
     '照常执行（不会「强制了 MODERN 却没人管 secret」）。')
h2('3.3　三档行为')
table([
    ['档位', 'advanced 变更', '其他', '启动日志'],
    ['SAFE（默认）', '无（严格上游默认）', '—', 'azureproxy.mode=SAFE (upstream defaults)'],
    ['ACCESS', 'log-command-executions = true', '—', 'azureproxy.mode=ACCESS applied (log-command-executions=true)'],
    ['EXP', 'log-command-executions = true（强制）announce-proxy-commands = true（强制）', 'player-info-forwarding-mode 未显式配置时设为 MODERN（显式则尊重）', 'azureproxy.mode=EXP applied (log-command-executions=true, announce-proxy-commands=true)'],
])
h2('3.4　T1：命令树注入修复')
para('**现象**：代理模式下 /server 红色未知命令、零 tab 补全。**根因**：EXP preset 残留下游改写 '
     'announce-proxy-commands=false——Velocity 依此不把代理命令树并入发给后端的 AvailableCommands，'
     '代理命令从客户端命令面消失。**修复**：EXP 分支强制 set(true)，启动日志同步打出该值。')
para('**证据（实测）**：')
table([
    ['验证项', '证据', '结论'],
    ['代理侧合并', 'AvailableCommands children 26 → 28（注入 server 子树 + velocity:callback）', 'PASS'],
    ['客户端命令包', 'pid=0x10 len=718，含 velocity:callback、server（+action/target）节点', 'PASS'],
    ['真实客户端（XY_TianQ）', '/ser tab 补全 server；命令文本白色（非未知红色）', 'PASS'],
    ['命令执行', '/server exp7 / /server arena 均正常', 'PASS'],
])

# ---------------- 四 ----------------
h1('四、26.1.2 协议对齐')
h2('4.1　版本矩阵')
table([
    ['协议号', '版本名', '状态'],
    ['775', 'MINECRAFT_26_1（26.1 / 26.1.1 / 26.1.2）', '已对齐（velocity 上游 26_1 表正确，无需修补）'],
    ['776', 'MINECRAFT_26_2（26.2）', '暂缓（未验证）'],
])
h2('4.2　StateRegistry 关键 id（实测）')
table([
    ['方向', '包', 'id', '备注'],
    ['CB', 'BundleDelimiter', '0x00', '1.19.4+ 常驻 0x00（偏移陷阱根源，见 4.3）'],
    ['CB', 'KeepAlive', '0x2C', '26_1 条目显式存在'],
    ['CB', 'JoinGame（LOGIN）', '0x31', '26_1 条目显式存在'],
    ['CB', 'SystemChat', '0x79', '26_1 条目显式存在'],
    ['CB', 'Respawn / Transfer / StoreCookie', '0x52 / 0x81 / 0x78', '26_1 条目显式存在'],
    ['CB', 'AvailableCommands', '0x10', '26_1 无独立条目，按 1.21.5 继承（客户端实测 pid=0x10）'],
    ['SB', 'PlayerLoaded', '0x2C', '26_1 条目显式存在（进世界必需）'],
    ['SB', 'ClientTickEnd / chat_command_signed', '0x0D / 0x08', 'StateRegistry 未注册（Velocity 透传），合成客户端直发、后端接受'],
])
h2('4.3　索引抽取 +1 假象（方法学教训）')
para('26.1 客户端注册表第 0 项是 BundleDelimiter（0x00）；按「顺序数索引」从数据表抽取 id 会被整体'
     ' +1 偏移骗过——早期曾据此列出的 27 处 velocity 表「修正」经真实客户端核验全部是误报，已 revert。'
     '结论：id 必须**真实客户端实测 + 合成客户端互证**，数据表与实测一致才可信。')
h2('4.4　26.2（776）')
para('未验证。后续 P0：26_2 表刷新（增量抽取）、mcclient 776 回归（握手/配置/PLAY 全程）、'
     '真实客户端 26.2 会话确认。')

# ---------------- 五 ----------------
h1('五、验证矩阵（实测）')
h2('5.1　服务器切换课题')
table([
    ['验证项', '结果'],
    ['/server arena 切换', '✅ 世界差异可辨（seed 987654321 / 地形 / 出生点）'],
    ['/server exp7 切回', '✅'],
    ['世界隔离', '✅ /say 只广播本世界'],
    ['转发身份一致性', '✅ 跨后端同 UUID 2b47bbd5-9532-3390-b1b6-8392740fa849（MODERN 转发）'],
    ['Op 独立性', '✅ 每后端各自 ops.json（Velocity 语义）'],
    ['命令面', '✅ T1 修复后 /server tab 补全 + 白色（§3.4）'],
])
h2('5.2　Fallback（后端宕机瞬时切换）')
para('kill arena → 代理日志**同一秒**出现 arena has disconnected 与 exp7 has connected；exp7 后端'
     ' joined the game；用户确认无世界加载等待（瞬时）。')
h2('5.3　E2E 协议链路')
para('mcclient 全流程：握手（proto=775）→ SetCompression → LoginSuccess → LoginAcknowledged → '
     '配置阶段（FinishedUpdate + KnownPacks/ClientInformation 应答）→ PLAY（PlayerLoaded + '
     'ClientTickEnd 保持「在世界上」）→ chat_command_signed → SystemChat 回显，VERDICT 全真；'
     '真实客户端（XY_TianQ）/say 信道确认 + T1 命令面确认。')

# ---------------- 六 ----------------
h1('六、测试工具链与测试布局')
h2('6.1　mcclient（合成 775 客户端）')
para('mcclient.py <host> <port> <name> [cmd]：离线登录 + 全程抓 raw-dump.bin + pid 直方图，退出码 '
     '0/1 表示断言。阶段要点：')
table([
    ['阶段', '关键动作'],
    ['LOGIN', 'handshake(0x00) → LoginStart(name+16B UUID) → SetCompression(0x03) → LoginSuccess(0x02) → LoginAcknowledged(0x03)'],
    ['CONFIG', '收 FinishedUpdate(0x03) → ack；CookieRequest 忽略；发 ClientInformation(0x00)；收 KnownPacks(0x0E) → 回 SelectKnownPacks(0x07)（空）'],
    ['PLAY', 'JoinGame 后发 PlayerLoaded(0x2C)（必须）；每 0.2s ClientTickEnd(0x0D)（26.1 起必需）；发 chat_command_signed(0x08)（签名占位，后端接受）；扫 SystemChat 回显'],
])
h2('6.2　mcping / RCON 冒烟')
para('mcping.py：MC 1.7+ status 探测（默认 proto=775）。exp7proxy-smoke.py（AzureBranches 侧）：'
     '经 RCON（25576/25578）发 EXP 链命令并以 @e 选择器 + 空列表 ghost 断言（规避 26.1.2 '
     'summon UUID:[I;…] 随机 UUID 缺陷），[PASS]/[FAIL] 逐项输出。')
h2('6.3　测试布局')
table([
    ['组件', '端口', '位置'],
    ['AzureProxy（EXP）', '25571（bind 127.0.0.1）', 'proxyrun-test/velocity.toml：servers exp7=25570 / arena=25572；try=[exp7,arena]；[azureproxy] mode="EXP"'],
    ['exp7 后端（b_linear_v4）', '25570 / RCON 25576', 'AzureBranches exp7-test/（exp7-rcon.py）'],
    ['arena 后端（seed 987654321）', '25572 / RCON 25578', 'AzureBranches exp7-test2/（arena-rcon.py）'],
])

# ---------------- 七 ----------------
h1('七、已知限制与后续工作')
h2('7.1　已知限制')
para('1. 26.2（776）未验证（协议表刷新 + mcclient 776 回归未做）；')
para('2. 后端侧 summon UUID:[I;…] 在 26.1.2 生成随机 UUID（上游缺陷），冒烟脚本以 @e + 空表断言规避；')
para('3. mcclient 仅离线登录（不覆盖正版登录/加密）；')
para('4. fallback 测试期间 arena 后端需手动重启（测试脚本化未完成）；')
para('5. 性能：代理参数未建立压力基线（单机小规模验证为主）。')
h2('7.2　后续工作（优先级）')
para('P0：26.2（776）协议刷新与回归；真实负载压力基线（多后端 + 长时间混合流量）；')
para('P1：更多后端（>2）切换/回环验证；测试编排脚本化（kill/重启/断言一键化）；')
para('P2：预设档位的完整配置面文档化（每档位逐项列出 velocity.toml 差异）。')

# ---------------- 八 ----------------
h1('八、结论')
para('AzureProxy 以「仓库只持补丁、上游按 ref 构建、fail-fast 防漂移」的构建哲学落地了 EXP7 的'
     '代理侧配套：EXP 三档预设经 nightconfig 绑定前改写无损接入上游校验链；T1 以真实客户端字节级'
     '证据闭环（命令树合并 + 命令包内容 + 客户端行为三层互证）；26.1.2 协议对齐确认上游表正确，'
     '并以「Bundle 0x00 +1 假象」案例固化了「实测优先于索引抽取」的方法学。服务器切换/fallback/'
     '转发身份/世界隔离全部实测通过。技术文档（TECHNICAL.md）与本文互补：前者是随实现更新的'
     '原理向规范，本文是 26.1.2 阶段的工程记录。')
blank()
para('文档：AzureProxy-26.1.2-0001　　版本：4.1.0-SNAPSHOT　　日期：2026 年 08 月 23 日　　'
     '项目：AzureProxy (https://github.com/XCxyTianQ/AzureProxy)', size=10.5)

doc.save(OUT)
print('saved', OUT)
