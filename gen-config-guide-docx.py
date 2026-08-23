# -*- coding: utf-8 -*-
"""Generate the dual-stack config guide docx (AzureDoc series).

Style mirrors AzureBranches-26.1.2-EXP7.docx / AzureProxy-26.1.2-0001.docx:
centered title block, justified body 10.5pt Times New Roman, h1 14pt / h2 12pt
bold, Table Grid tables; code blocks rendered in Consolas.
Output: F:\\AzureCore\\AzureDoc\\AzureProxy-26.1.2-CONFIG-GUIDE.docx
"""
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = r'F:\AzureCore\AzureDoc\AzureProxy-26.1.2-CONFIG-GUIDE.docx'
FONT = 'Times New Roman'
MONO = 'Consolas'

doc = docx.Document()
st = doc.styles['Normal']
st.font.name = FONT
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def mkruns(p, text, bold=False):
    for chunk in text.split('**'):
        if not chunk:
            continue
        r = p.add_run(chunk)
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        r.font.size = Pt(10.5)
        r.bold = bold


def para(text, size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space=Pt(3)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space
    p.paragraph_format.line_spacing = 1.0
    mkruns(p, text, bold)
    for r in p.runs:
        r.font.size = Pt(size)
    return p


def code(text, size=9.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = MONO
    r._element.rPr.rFonts.set(qn('w:eastAsia'), MONO)
    r.font.size = Pt(size)
    return p


def h1(text):
    return para(text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space=Pt(8))


def h2(text):
    return para(text, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space=Pt(6))


def blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return p


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            c = t.cell(ri, ci)
            c.text = ''
            p = c.paragraphs[0]
            mkruns(p, cell, bold=(ri == 0))
            for r in p.runs:
                r.font.size = Pt(10.5)
    return t


# ---------------- title block ----------------
para('AzureProxy + AzureBranches', size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(2))
para('配置手册（简单版）', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(2))
para('从零搭起「代理 + 双后端」小服：能进、能切、计分板正常、/team 可用', size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(4))
para('配置基线 26.1.2-EXP7Plus / v26.1.2-AP-0001　　2026 年 08 月 23 日',
     size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(2))
para('基于 AzureBranches 26.1.2-EXP7Plus（后端）+ AzureProxy 26.1.2-AP-0001（代理）',
     size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(2))
para('配套：AzureBranches README / AzureProxy CONFIG-GUIDE.md / TECHNICAL.md（原理向）',
     size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space=Pt(6))
blank()

# ---------------- 0 ----------------
h1('零、这套东西是什么（30 秒版）')
code('玩家 ──► AzureProxy（代理，25571）──► 后端 exp7（25570，正式世界）')
code('                                  └──► 后端 arena（25572，测试世界）')
para('· **AzureBranches**＝服务器本体（Folia 下游）。它让单个后端拥有：b_linear_v4 存储引擎、命令 OCC（EXP 链）、计分板/团队修复等。')
para('· **AzureProxy**＝入口代理（Velocity 下游）。它做三件事：一个地址进多个后端、玩家随时切服（/server）、后端掉线自动换后备。')
para('· 玩家只连代理，永远不直接连后端。')

# ---------------- 1 ----------------
h1('一、准备')
table([
    ['需要', '说明'],
    ['JDK 25', '运行与构建都需要'],
    ['jar 文件', '方式一：GitHub Releases 下载；方式二：本地构建（见第 2 节）'],
    ['目录规划', '每个后端一个文件夹（如 server-exp7/、server-arena/），代理一个文件夹（如 proxy/）；互相独立，别混放'],
])

# ---------------- 2 ----------------
h1('二、拿 jar')
h2('2.1　方式一：下载（推荐）')
para('· 后端：AzureBranches Releases → v26.1.2-EXP7Plus → azurebranches-server-26.1.2-AB-0002-EXP7Plus.jar；')
para('· 代理：AzureProxy Releases → v26.1.2-AP-0001 → azureproxy-proxy-26.1.2-AP-0001.jar。')
h2('2.2　方式二：本地构建')
code('后端（AzureBranches 目录）')
code('    ./gradlew.bat :azurebranches-server:buildFolia :azurebranches-server:mergeJar')
code('    产物：folia-server/build/libs/azurebranches-server-26.1.2-AB-0002-EXP7Plus.jar')
code('')
code('代理（AzureProxy 目录）')
code('    ./gradlew buildAzureProxyJar')
code('    产物：build/libs/azureproxy-proxy-4.1.0-SNAPSHOT-all.jar')

# ---------------- 3 ----------------
h1('三、搭后端（每个服一次）')
para('以 server-exp7/ 为例（arena 同法，换个端口即可）。')
h2('3.1　基础文件')
code('mkdir server-exp7 && cd server-exp7')
code('# 放入 azurebranches-server-26.1.2-AB-0002-EXP7Plus.jar')
code('echo "eula=true" > eula.txt')
h2('3.2　server.properties（改 / 加这几行即可）')
code('server-port=25570')
code('online-mode=false')
code('level-name=world-exp7')
code('motd=My AzureBranches Server')
code('enable-rcon=true          # 可选：远程控制台')
code('rcon.port=25576')
code('rcon.password=改一个强密码')
h2('3.3　config/paper-global.yml → 打开代理转发（关键！）')
code('velocity:')
code('  enabled: true')
code('  online-mode: false')
code('  secret: 一个和代理完全一样的随机串')
para('**警告**：secret 必须与代理的 forwarding.secret 文件内容**逐字符一致**，否则玩家身份/UUID 会错乱或登录被拒。建议用 24+ 位随机串。')
h2('3.4　azurebranches_global_config.toml（服务器根目录）→ 开启 EXP 档与存储引擎（可选）')
code('[command_blocks]')
code('mode = "EXP"          # SAFE（上游默认）| ACCESS | EXP；EXP = 命令链 OCC 全套')
code('')
code('[storage]')
code('region_format = "b_linear_v4"   # "mca"（默认，原版）| "b_linear_v4"')
code('')
code('[storage.linear]')
code('compression_level = 1            # 1..22，压缩级别，默认 1')
para('建议：新世界直接用 b_linear_v4（写入更快 + 四层校验）；老世界（已有 MCA 数据）先备份再考虑。')
h2('3.5　启动')
code('java -Xmx2G -jar azurebranches-server-26.1.2-AB-0002-EXP7Plus.jar nogui')
para('看到 Done (…) 即成功。三个维度 + entities（r.*.mca）落在该世界区域里，v4 主文件与 .swp 交换文件同目录。')
para('**注意**：开了 velocity.enabled: true 的后端，只能经代理进入——玩家直连会被拒（转发校验）。这是正常现象。')

# ---------------- 4 ----------------
h1('四、搭代理')
para('在 proxy/ 目录：放入代理 jar + forwarding.secret 文件：')
code('proxy/')
code('├── azureproxy-proxy-26.1.2-AP-0001.jar')
code('├── forwarding.secret        # 内容 = 后端 paper-global.yml 的 velocity.secret（一模一样）')
code('└── velocity.toml            # 见下')
h2('4.1　velocity.toml（最小可跑样例）')
code('bind = "0.0.0.0:25571"          # 玩家连的地址；本机测试可写 127.0.0.1')
code('online-mode = false')
code('player-info-forwarding-mode = "MODERN"')
code('forwarding-secret-file = "forwarding.secret"')
code('')
code('[servers]')
code('exp7 = "127.0.0.1:25570"')
code('arena = "127.0.0.1:25572"')
code('')
code('try = ["exp7", "arena"]          # 进服与掉线时按顺序尝试')
code('')
code('[azureproxy]')
code('mode = "EXP"                     # SAFE | ACCESS | EXP，见第 5 节')
h2('4.2　启动')
code('java -jar azureproxy-proxy-26.1.2-AP-0001.jar')
para('看到 Booting up AzureProxy … 与 [AzureProxy] azureproxy.mode=EXP applied (log-command-executions=true, announce-proxy-commands=true) 即成功。')

# ---------------- 5 ----------------
h1('五、代理档位怎么选')
table([
    ['mode', '一句话', '适用'],
    ['SAFE', '和原版 Velocity 一模一样，零改动', '只想当纯代理'],
    ['ACCESS', '多记录命令执行日志', '排错 / 观察'],
    ['EXP', '观察 + 强制代理命令树（/server 可 tab 补全）+ 未配置时强制 MODERN 转发', '常规用途（推荐）'],
])
para('**为什么日常就用 EXP**：T1 修复后，EXP 档保证客户端命令列表里 /server 是白色、可补全；改成 SAFE 只影响这些预设，不影响进服。')

# ---------------- 6 ----------------
h1('六、后端档位与代理档位对照')
table([
    ['后端 command_blocks.mode', '代理 [azureproxy] mode', '效果'],
    ['SAFE', 'SAFE', '全上游默认，最保守'],
    ['EXP', 'EXP │ ACCESS', '命令链 OCC + 代理命令面（推荐组合）'],
])
para('两个开关独立：代理档管「网络/命令面」，后端档管「命令链 OCC」。')

# ---------------- 7 ----------------
h1('七、日常使用')
table([
    ['想干什么', '怎么做'],
    ['进服', '客户端连 你的IP:25571（或服务器列表地址），自动进 try 的第一个'],
    ['切服', '游戏内 /server exp7 或 /server arena；输入 /server 有自动补全'],
    ['世界隔离', '每个后端是独立世界：/say、/tell 只在本服广播'],
    ['管理员', '每后端各自 ops.json（op 玩家名 分别执行），互不影响'],
    ['后端宕机', '代理自动切换 try 里下一个可用后端，通常同一秒完成'],
    ['计分板', 'EXP7Plus 已修复：重进后 sidebar 正常（会重发 SetObjective/Display/Score）'],
    ['团队', 'EXP7Plus 已恢复：/team add|join|modify color|list … 全部可用'],
])

# ---------------- 8 ----------------
h1('八、常见问题排查')
table([
    ['现象', '原因与解法'],
    ['重进后计分板消失（服务端数据还在）', '后端不是 EXP7Plus 版本 → 换 azurebranches-server-26.1.2-AB-0002-EXP7Plus.jar'],
    ['/team 报 Unknown command', '同上（旧版本 Folia 禁用了 team，EXP7Plus 已恢复）'],
    ['/server 红色未知命令 / 无 tab 补全', '代理 [azureproxy] mode 不是 EXP（或高级配置 announce-proxy-commands=false）→ 用 EXP 档'],
    ['能直连后端、经代理进不去 / 被拒', '后端开了 velocity.enabled: true → 必须走代理；或 secret 不一致'],
    ['UUID 每次不同 / 身份错乱', 'velocity.secret（后端）≠ forwarding.secret（代理）→ 改成一致'],
    ['Booting up AzureProxy 未见 EXP 行', '检查 velocity.toml 的 [azureproxy] mode 段拼写 / 位置（应在文件末尾独立一节）'],
    ['切服卡在加载', '目标后端没启动或端口错 → 先确认后端 Done 且端口匹配 [servers]'],
])

# ---------------- 9 ----------------
h1('九、版本对应关系')
table([
    ['组件', '版本', 'jar'],
    ['后端', 'v26.1.2-EXP7Plus（当前）', 'azurebranches-server-26.1.2-AB-0002-EXP7Plus.jar'],
    ['代理', 'v26.1.2-AP-0001', 'azureproxy-proxy-26.1.2-AP-0001.jar'],
])
para('两端可以分开选版（例如后端 EXP7、代理 AP-0001），但 EXP7Plus 的两个修复（计分板登录同步、/team）只在后端侧，与代理版本无关。')
blank()
para('文档：AzureProxy-26.1.2-CONFIG-GUIDE　　版本：26.1.2-EXP7Plus / v26.1.2-AP-0001　　日期：2026 年 08 月 23 日　　项目：AzureProxy + AzureBranches (https://github.com/XCxyTianQ/AzureProxy)', size=10.5)

doc.save(OUT)
print('saved', OUT)
